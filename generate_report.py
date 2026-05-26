#!/usr/bin/env python3
"""生成《软件构造》实验报告 Word 文档"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ============ 全局样式设置 ============
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5

# 设置页边距
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)


def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading_styled(text, level=1):
    """添加带样式的标题"""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        if level == 1:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0, 0, 0)
        elif level == 2:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 0, 0)
        elif level == 3:
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_para(text, bold=False, font_size=None, align=None, font_name=None, space_after=Pt(6)):
    """添加段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if font_size:
        run.font.size = font_size
    if font_name:
        run.font.name = font_name
        run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if align:
        p.alignment = align
    p.paragraph_format.space_after = space_after
    return p


def add_table(headers, rows, col_widths=None):
    """添加表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, 'D9E2F3')

    # 数据行
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(10.5)
            run.font.name = '宋体'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 设置列宽
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # 表后空行
    return table


def add_code_block(lines):
    """添加代码块（灰色背景、等宽字体）"""
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        # 灰色背景
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'F2F2F2')
        shading.set(qn('w:val'), 'clear')
        run.element.rPr.append(shading)


# ============ 封面 ============
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('《软件构造》实验报告')
run.bold = True
run.font.size = Pt(26)
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

doc.add_paragraph()
doc.add_paragraph()

# 封面信息表格
info_data = [
    ['实验名称', '实验一：AI 辅助需求建模与系统骨架构造'],
    ['院    系', '计算机科学与工程学院'],
    ['班    级', '23040301'],
    ['姓    名', '刘禹赫'],
    ['学    号', '20233063'],
]
info_table = doc.add_table(rows=len(info_data), cols=2)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (k, v) in enumerate(info_data):
    c0 = info_table.rows[i].cells[0]
    c1 = info_table.rows[i].cells[1]
    c0.text = ''
    c1.text = ''
    r0 = c0.paragraphs[0].add_run(k)
    r0.bold = True
    r0.font.size = Pt(14)
    r0.font.name = '黑体'
    r0.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    c0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = c1.paragraphs[0].add_run(v)
    r1.font.size = Pt(14)
    r1.font.name = '宋体'
    r1.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    c0.width = Cm(3)
    c1.width = Cm(8)

# 去掉表格边框（封面表格不显示线）
for row in info_table.rows:
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        borders = OxmlElement('w:tcBorders')
        for edge in ['top', 'left', 'bottom', 'right']:
            el = OxmlElement(f'w:{edge}')
            el.set(qn('w:val'), 'none')
            el.set(qn('w:sz'), '0')
            borders.append(el)
        tcPr.append(borders)

doc.add_page_break()

# ============ 正文 ============

# 一、实验目标
add_heading_styled('一、实验目标', level=1)
goals = [
    '掌握将自然语言需求转化为用户故事、功能需求、非功能需求、验收标准的需求工程方法。',
    '借助 AI 完成领域建模、模块划分、项目结构设计。',
    '完成项目初始化，搭建符合 Spring MVC 分层架构的工程骨架。',
    '建立需求基线与架构基线，为后续开发提供依据。',
]
for i, g in enumerate(goals, 1):
    add_para(f'{i}. {g}')

# 二、实验内容
add_heading_styled('二、实验内容', level=1)
contents = [
    ('需求梳理', '明确系统角色、核心场景、功能边界。'),
    ('领域建模', '识别实体、值对象、业务规则与不变量。'),
    ('架构设计', '采用 Spring Boot + Spring MVC 分层架构。'),
    ('工程搭建', '初始化 Maven 项目，生成标准目录结构。'),
    ('AI 辅助', '使用 Prompt 驱动 AI 输出初稿并人工修订。'),
]
for i, (title, desc) in enumerate(contents, 1):
    add_para(f'{i}. {title}：{desc}')

# 三、项目概述
add_heading_styled('三、项目概述', level=1)
add_para(
    'StudyFlow 是面向学生的学习任务管理系统，用于任务规划、优先级管理、'
    '到期提醒与学习进度统计，解决任务混乱、优先级不清、截止时间遗漏、进度无统计等问题。'
)
add_para('核心功能：', bold=True)
funcs = ['任务增删改查', '优先级 / 状态管理', '提醒策略配置', '任务排序与筛选', '学习完成率统计']
for f in funcs:
    p = doc.add_paragraph(f, style='List Bullet')
    for run in p.runs:
        run.font.name = '宋体'
        run.font.size = Pt(12)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 四、系统架构与项目结构
add_heading_styled('四、系统架构与项目结构', level=1)

add_heading_styled('1. 架构模式', level=2)
add_para('采用 Spring MVC 分层架构：Controller → Service → Mapper → DB，各层职责单一、低耦合。')

add_heading_styled('2. 标准项目结构', level=2)
tree_lines = [
    'StudyFlow/',
    '├── src/main/java/com/studyflow/',
    '│   ├── controller/          # 请求接收与响应',
    '│   ├── service/             # 业务逻辑',
    '│   │   └── impl/',
    '│   ├── mapper/              # 数据库访问',
    '│   ├── entity/              # 数据实体',
    '│   ├── dto/                 # 请求参数（预留）',
    '│   ├── vo/                  # 返回结果（预留）',
    '│   ├── config/              # 配置（预留）',
    '│   ├── exception/           # 异常处理（预留）',
    '│   ├── constant/            # 常量（预留）',
    '│   ├── utils/               # 工具（预留）',
    '│   └── StudyFlowApplication.java',
    '├── src/main/resources/',
    '│   ├── mapper/              # MyBatis XML 映射',
    '│   │   └── TaskMapper.xml',
    '│   ├── sql/',
    '│   │   └── schema.sql       # 数据库建表脚本',
    '│   └── application.yml      # Spring Boot 配置',
    '├── src/test/java/',
    '├── pom.xml',
    '└── README.md',
]
add_code_block(tree_lines)
doc.add_paragraph()

add_heading_styled('3. 技术栈', level=2)
add_para('Java 8 + Spring Boot 2.7.18 + MyBatis 2.3.2 + MySQL 8.x + Maven 3.x')

# 五、实验结果与验证
add_heading_styled('五、实验结果与验证', level=1)
results = [
    '完成需求文档，覆盖所有功能与约束。',
    '完成领域模型，实体关系清晰。',
    '搭建可运行 Spring Boot 工程骨架，分层职责明确。',
    '业务不变量已在需求中固化，可用于编码校验。',
]
for i, r in enumerate(results, 1):
    add_para(f'{i}. {r}')

# 六、遇到的问题与解决
add_heading_styled('六、遇到的问题与解决', level=1)
add_table(
    ['问题', '解决方案'],
    [
        ['AI 输出领域分层（domain/application）与 Spring MVC 冲突', '改为 controller/service/mapper 标准结构'],
        ['AI 建议本地文件存储', '改为 MySQL + MyBatis，保证数据持久化'],
        ['AI 推荐 Java 17', '统一为 Java 8，符合课程要求'],
    ],
    col_widths=[7, 7],
)

# 七、个人反思
add_heading_styled('七、个人反思', level=1)
reflections = [
    'AI 能快速生成需求草案、结构模板、代码骨架，大幅提升效率。',
    '工程决策必须人工把控：技术选型、架构取舍、业务规则、规范对齐。',
    '需求建模是软件开发的核心，好的模型能减少后期返工。',
    '本次实验掌握了需求工程、领域建模、分层架构、AI 辅助开发全流程，为后续实验打下基础。',
]
for i, r in enumerate(reflections, 1):
    add_para(f'{i}. {r}')

# ============ 分页 ============
doc.add_page_break()

# ============ 附录一：需求说明文档 ============
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run('附录一：需求说明文档')
run.bold = True
run.font.size = Pt(18)
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
doc.add_paragraph()

add_heading_styled('1. 用户角色', level=2)
add_table(
    ['角色', '描述'],
    [['学生用户（Student）', '系统唯一角色，负责管理个人学习任务']],
    col_widths=[4, 10],
)

add_heading_styled('2. 用户故事', level=2)
add_table(
    ['编号', '用户故事', '验收标准'],
    [
        ['US-01', '作为学生，我希望创建一条学习任务，包含标题、描述和截止时间', '标题不能为空；截止时间不得早于当前时间；创建成功返回任务详情'],
        ['US-02', '作为学生，我希望为任务设置高/中/低优先级', '默认优先级为 MEDIUM；高优先级必须设置截止时间'],
        ['US-03', '作为学生，我希望为任务配置提醒策略', '支持 DEADLINE_ONE_DAY、HIGH_PRIORITY_IMMEDIATE、DAILY_UNFINISHED_SUMMARY'],
        ['US-04', '作为学生，我希望修改任务信息', '已完成状态的任务不允许修改'],
        ['US-05', '作为学生，我希望删除任务', '删除后任务不可恢复、不可访问'],
        ['US-06', '作为学生，我希望更新任务状态', '状态只能 TODO → IN_PROGRESS → COMPLETED，不可回退'],
        ['US-07', '作为学生，我希望按优先级查看任务', '高优先级排在前面；同等优先级按截止时间升序'],
        ['US-08', '作为学生，我希望查看学习统计', '显示总任务数、已完成数、未完成数、完成率'],
    ],
    col_widths=[1.5, 5, 7.5],
)

add_heading_styled('3. 功能需求', level=2)
add_table(
    ['编号', '功能', '说明'],
    [
        ['FR-01', '创建任务', '输入标题、描述、截止时间、优先级，创建新任务'],
        ['FR-02', '修改任务', '根据 ID 修改任务信息（标题、描述、截止时间、优先级）'],
        ['FR-03', '删除任务', '根据 ID 删除任务，物理删除'],
        ['FR-04', '设置优先级', '支持 HIGH、MEDIUM、LOW 三档'],
        ['FR-05', '更新状态', '支持 TODO、IN_PROGRESS、COMPLETED，单向流转'],
        ['FR-06', '排序查询', '默认按优先级（HIGH > MEDIUM > LOW）+ 截止时间升序排列'],
        ['FR-07', '提醒策略管理', '为每个任务配置提醒策略（截止前一天、高优立即提醒、每日未完成汇总）'],
        ['FR-08', '学习进度统计', '返回总任务数、已完成数、未完成数、完成率'],
        ['FR-09', '任务筛选', '支持按状态、优先级筛选，支持组合筛选'],
        ['FR-10', '到期任务查询', '查询 24 小时内到期的任务（预留接口）'],
    ],
    col_widths=[1.5, 3.5, 9],
)

add_heading_styled('4. 非功能需求', level=2)
add_table(
    ['编号', '类别', '说明'],
    [
        ['NFR-01', '技术栈', 'Java 8 + Spring Boot 2.x + MyBatis 3.x + MySQL 8.x + Maven 3.x'],
        ['NFR-02', '架构', 'Spring MVC 分层（Controller / Service / Mapper / Entity）'],
        ['NFR-03', '持久化', 'MySQL 存储，数据不丢失'],
        ['NFR-04', '可测试', 'Service 层支持单元测试与 Mock'],
        ['NFR-05', '可维护', '单一职责、低耦合、规范命名'],
        ['NFR-06', '可扩展', '提醒策略支持策略模式扩展'],
    ],
    col_widths=[1.5, 3, 9.5],
)

add_heading_styled('5. 业务不变量', level=2)
invariants = [
    '截止时间 ≥ 创建时间。',
    'HIGH 优先级任务必须设置截止时间。',
    '状态只能按 TODO → IN_PROGRESS → COMPLETED 单向流转，不可回退。',
    '任务标题不能为空。',
    '删除后任务不可访问。',
]
for i, inv in enumerate(invariants, 1):
    add_para(f'{i}. {inv}')

add_heading_styled('6. 验收标准', level=2)
add_para('覆盖以下场景的正确 / 异常流程：')
scenarios = [
    '创建任务：正常创建、标题为空、高优无截止时间、截止时间已过',
    '更新任务：正常修改、修改已完成任务',
    '删除任务：正常删除、删除不存在任务',
    '状态流转：正常流转、非法回退',
    '查询统计：有任务、无任务时完成率为 0',
]
for s in scenarios:
    p = doc.add_paragraph(s, style='List Bullet')
    for run in p.runs:
        run.font.name = '宋体'
        run.font.size = Pt(12)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ============ 分页 ============
doc.add_page_break()

# ============ 附录二：领域模型 ============
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run('附录二：领域模型（文字建模说明）')
run.bold = True
run.font.size = Pt(18)
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
doc.add_paragraph()

add_heading_styled('1. 核心实体', level=2)

add_heading_styled('Task（学习任务）', level=3)
add_table(
    ['属性', '类型', '说明'],
    [
        ['id', 'Long', '主键，自增'],
        ['title', 'String', '任务标题（必填，≤100 字符）'],
        ['description', 'String', '任务描述（可选）'],
        ['priority', 'Priority', '优先级（值对象）'],
        ['status', 'TaskStatus', '状态（值对象）'],
        ['deadline', 'Date', '截止时间'],
        ['reminderPolicy', 'ReminderPolicy', '提醒策略（值对象）'],
        ['createTime', 'Date', '创建时间（系统自动）'],
        ['updateTime', 'Date', '更新时间（系统自动）'],
    ],
    col_widths=[3, 3.5, 7.5],
)

add_heading_styled('ProgressReport（学习统计报表）', level=3)
add_table(
    ['属性', '类型', '说明'],
    [
        ['totalCount', 'int', '总任务数'],
        ['completedCount', 'int', '已完成任务数'],
        ['unfinishedCount', 'int', '未完成任务数'],
        ['completionRate', 'double', '完成率（百分比）'],
    ],
    col_widths=[4, 3, 7],
)
add_para('ProgressReport 不是独立持久化实体，而是基于 Task 数据实时聚合生成。')

add_heading_styled('2. 值对象', level=2)

add_para('Priority（优先级）', bold=True)
add_table(
    ['值', '含义'],
    [
        ['HIGH', '高优先级，必须设置截止时间'],
        ['MEDIUM', '中优先级（默认值）'],
        ['LOW', '低优先级'],
    ],
    col_widths=[3, 11],
)

add_para('TaskStatus（任务状态）', bold=True)
add_table(
    ['值', '含义'],
    [
        ['TODO', '未开始（默认值）'],
        ['IN_PROGRESS', '进行中'],
        ['COMPLETED', '已完成'],
    ],
    col_widths=[4, 10],
)
add_para('状态流转规则：TODO → IN_PROGRESS → COMPLETED（单向，不可回退）')

add_para('ReminderPolicy（提醒策略）', bold=True)
add_table(
    ['值', '含义'],
    [
        ['DEADLINE_ONE_DAY', '截止时间前一天提醒'],
        ['HIGH_PRIORITY_IMMEDIATE', '高优先级任务立即提醒'],
        ['DAILY_UNFINISHED_SUMMARY', '每日未完成任务汇总提醒'],
    ],
    col_widths=[5, 9],
)

add_heading_styled('3. 实体关系', level=2)
# 用文字描述关系
rel_lines = [
    '┌─────────────────────────────────────────────┐',
    '│                   Task                       │',
    '│─────────────────────────────────────────────│',
    '│  id: Long (PK)                              │',
    '│  title: String                              │',
    '│  description: String                        │',
    '│  priority: Priority ──────────────┐         │',
    '│  status: TaskStatus ────────────┐ │         │',
    '│  deadline: Date                 │ │         │',
    '│  reminderPolicy: ReminderPolicy │ │         │',
    '│  createTime: Date               │ │         │',
    '│  updateTime: Date               │ │         │',
    '└──────────────────────────────────┼─┼─────────┘',
    '                                   │ │',
    '         包含（has-a）             │ │',
    '    ┌──────────────┐              │ │',
    '    │   Priority   │              │ │',
    '    │──────────────│              │ │',
    '    │  HIGH        │◄─────────────┘ │',
    '    │  MEDIUM      │                │',
    '    │  LOW         │                │',
    '    └──────────────┘                │',
    '                                    │',
    '    ┌──────────────┐                │',
    '    │  TaskStatus  │                │',
    '    │──────────────│                │',
    '    │  TODO        │◄───────────────┘',
    '    │  IN_PROGRESS │',
    '    │  COMPLETED   │',
    '    └──────────────┘',
    '',
    '    ┌──────────────────┐',
    '    │ ReminderPolicy   │',
    '    │──────────────────│',
    '    │  DEADLINE_ONE_DAY│',
    '    │  HIGH_PRIORITY_  │',
    '    │    IMMEDIATE     │',
    '    │  DAILY_UNFINISHED│',
    '    │    _SUMMARY      │',
    '    └──────────────────┘',
    '',
    '┌────────────────────────────────────────┐',
    '│          ProgressReport                │',
    '│────────────────────────────────────────│',
    '│  totalCount: int                       │',
    '│  completedCount: int                   │',
    '│  unfinishedCount: int                  │',
    '│  completionRate: double                │',
    '└────────────────────────────────────────┘',
    '         ▲',
    '         │ 基于（derived from）',
    '    Task 集合',
]
add_code_block(rel_lines)
doc.add_paragraph()

add_heading_styled('4. 业务规则（不变量）', level=2)
add_table(
    ['编号', '规则', '实现位置'],
    [
        ['IR-01', '任务标题不能为空', 'Service 层 createTask 校验'],
        ['IR-02', '截止时间不能早于创建时间', 'Service 层 createTask 校验'],
        ['IR-03', 'HIGH 优先级必须设置截止时间', 'Service 层 createTask 校验'],
        ['IR-04', '状态只能 TODO → IN_PROGRESS → COMPLETED', 'Service 层 updateTaskStatus 校验'],
        ['IR-05', '已完成任务不允许修改', 'Service 层 updateTask 校验'],
        ['IR-06', '删除后任务不可访问', '物理删除，selectById 返回 null'],
    ],
    col_widths=[1.5, 6, 6.5],
)

# ============ 分页 ============
doc.add_page_break()

# ============ 附录三：AI Prompt 记录表 ============
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run('附录三：AI Prompt 记录表')
run.bold = True
run.font.size = Pt(18)
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
doc.add_paragraph()

add_table(
    ['序号', 'Prompt 内容', 'AI 输出（摘要）', '人工修订'],
    [
        ['1', '补充完善需求清单：角色、用户故事、功能需求、非功能需求、验收标准', '输出完整需求初稿，包含用户故事、功能/非功能需求、验收标准', '统一技术栈为 Java 8 + MySQL；修正分层为 Spring MVC（而非 domain/application）'],
        ['2', '基于需求做领域建模：实体、值对象、关系、不变量', '输出 Task 实体、策略对象、状态对象、业务规则', '补充状态不可逆约束；补充高优先级必须带截止时间的校验'],
        ['3', '用 Java 8 + Spring MVC + MySQL 设计项目结构', '输出通用分层目录结构', '调整为标准 Spring Boot 结构：controller / service / mapper / entity；添加 dto / vo / exception 等预留包'],
        ['4', '生成 StudyFlow 项目骨架代码', '输出 pom.xml、application.yml、实体类、Mapper、Service、Controller', '调整日期处理为 java.util.Date；补充 MyBatis XML 映射文件；补充数据库建表 SQL'],
        ['5', '实验报告结构与内容', '输出报告初稿', '结合实际项目成果重写实验结果与反思部分'],
    ],
    col_widths=[1.2, 4, 4.3, 4.5],
)

# ============ 页脚说明 ============
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— 全文完 —')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(128, 128, 128)

# ============ 保存 ============
output_path = '/Users/lyh/testclass/Software Construction/StudyFlow/实验报告.docx'
doc.save(output_path)
print(f'Word 文档已生成：{output_path}')
