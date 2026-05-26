#!/usr/bin/env python3
"""生成《软件构造》实验二 - 接口签名与行为约束 Word 文档"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ============ 全局样式设置 ============
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)


def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        if level == 1:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0, 0, 0)
        elif level == 2:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 0, 0)
        elif level == 3:
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_para(text, bold=False, font_size=None, align=None, font_name=None, space_after=Pt(6)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if font_size:
        run.font.size = font_size
    if font_name:
        run.font.name = font_name
        run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if align:
        p.alignment = align
    p.paragraph_format.space_after = space_after
    return p


def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, 'D9E2F3')

    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(10.5)
            run.font.name = '宋体'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return table


def add_code_block(lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'F2F2F2')
        shading.set(qn('w:val'), 'clear')
        run.element.rPr.append(shading)


# ============ 封面 ============
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('《软件构造》实验报告')
run.bold = True
run.font.size = Pt(26)
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

doc.add_paragraph()
doc.add_paragraph()

info_data = [
    ['实验名称', '实验二：契约驱动实现与测试优先构造'],
    ['院    系', '计算机科学与工程学院'],
    ['班    级', '23040301'],
    ['姓    名', '刘禹赫'],
    ['学    号', '20233063'],
]
info_table = doc.add_table(rows=len(info_data), cols=2)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (k, v) in enumerate(info_data):
    c0 = info_table.rows[i].cells[0]
    c1 = info_table.rows[i].cells[1]
    c0.text = ''
    c1.text = ''
    r0 = c0.paragraphs[0].add_run(k)
    r0.bold = True
    r0.font.size = Pt(14)
    r0.font.name = '黑体'
    r0.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    c0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = c1.paragraphs[0].add_run(v)
    r1.font.size = Pt(14)
    r1.font.name = '宋体'
    r1.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    c0.width = Cm(3)
    c1.width = Cm(8)

for row in info_table.rows:
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        borders = OxmlElement('w:tcBorders')
        for edge in ['top', 'left', 'bottom', 'right']:
            el = OxmlElement(f'w:{edge}')
            el.set(qn('w:val'), 'none')
            el.set(qn('w:sz'), '0')
            borders.append(el)
        tcPr.append(borders)

doc.add_page_break()

# ============ 正文 ============

# 一、实验目标
add_heading_styled('一、实验目标', level=1)
goals = [
    '学会将需求转化为接口、契约和测试用例。',
    '理解测试优先与小步提交的构造方式。',
    '学会利用 AI 生成测试样例、边界条件和实现候选。',
    '掌握"先写测试，再让 AI 生成实现"的实践流程。',
]
for i, g in enumerate(goals, 1):
    add_para(f'{i}. {g}')

# 二、实验内容
add_heading_styled('二、实验内容', level=1)
add_para(
    '在实验一工程基础上，实现 StudyFlow 的核心任务管理能力。'
    '本次实验先给出接口签名和行为约束（前置条件、后置条件、异常条件、边界情况），'
    '再由 AI 补全测试用例，最后完成实现。'
)
add_para('核心功能覆盖：', bold=True)
funcs = ['创建任务', '更新任务状态', '按优先级排序', '查询到期任务', '按用户统计任务完成情况']
for f in funcs:
    p = doc.add_paragraph(f, style='List Bullet')
    for run in p.runs:
        run.font.name = '宋体'
        run.font.size = Pt(12)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 三、接口定义
add_heading_styled('三、接口签名定义', level=1)

add_para(
    '根据实验一的需求分析和领域模型，我梳理了 TaskService 接口。'
    '以下是接口签名以及每个方法的前置条件、后置条件、异常条件和边界情况。'
)

add_heading_styled('1. 实体定义', level=2)
entity_code = [
    'public class Task {',
    '    private Long id;              // 主键，自增',
    '    private String title;         // 任务标题（必填）',
    '    private String description;   // 任务描述（可选）',
    '    private String priority;      // 优先级：HIGH / MEDIUM / LOW',
    '    private String status;        // 状态：TODO / IN_PROGRESS / COMPLETED',
    '    private Date deadline;        // 截止时间',
    '    private String reminderPolicy;// 提醒策略',
    '    private Date createTime;      // 创建时间',
    '    private Date updateTime;      // 更新时间',
    '}',
]
add_code_block(entity_code)
doc.add_paragraph()

add_heading_styled('2. TaskService 接口签名', level=2)
interface_code = [
    'public interface TaskService {',
    '    Task createTask(Task task);',
    '    Task updateTask(Long id, Task task);',
    '    void deleteTask(Long id);',
    '    Task getTaskById(Long id);',
    '    List<Task> getAllTasks();',
    '    List<Task> getTasksByStatus(String status);',
    '    List<Task> getTasksByPriority(String priority);',
    '    void updateTaskStatus(Long id, String status);',
    '    int getTotalCount();',
    '    int getCompletedCount();',
    '    int getUnfinishedCount();',
    '    double getCompletionRate();',
    '}',
]
add_code_block(interface_code)
doc.add_paragraph()

# 四、行为约束
add_heading_styled('四、各方法行为约束', level=1)

# ---------- 方法 1: createTask ----------
add_heading_styled('1. Task createTask(Task task) — 创建任务', level=2)

add_table(
    ['维度', '描述'],
    [
        ['前置条件', 'task 不为 null；task.title 不为 null 且非空字符串（去除首尾空格后长度 > 0）'],
        ['后置条件', '返回的 Task 对象包含系统分配的 id；若未指定 status，默认设为 "TODO"；数据库中新增一条记录；createTime 由系统自动填充'],
        ['异常条件',
         '(1) task 为 null → NullPointerException\n'
         '(2) title 为空或纯空格 → IllegalArgumentException("任务标题不能为空")\n'
         '(3) priority 为 "HIGH" 但 deadline 为 null → IllegalArgumentException("高优先级任务必须设置截止时间")\n'
         '(4) deadline 早于当前时间 → IllegalArgumentException("截止时间不能早于当前时间")'],
        ['边界情况',
         '(1) title 为纯空格 "   " → 应视为无效，抛出异常\n'
         '(2) priority 为 "MEDIUM" 或 "LOW" 时，不要求设置 deadline\n'
         '(3) deadline 恰好等于当前时间 → 理论上无法精确构造，但若传入过去时间应拒绝\n'
         '(4) title 为正常字符串（如 "复习高数"） → 正常创建'],
    ],
    col_widths=[2.5, 11.5],
)

# ---------- 方法 2: updateTask ----------
add_heading_styled('2. Task updateTask(Long id, Task task) — 更新任务', level=2)

add_table(
    ['维度', '描述'],
    [
        ['前置条件', 'id 不为 null；task 不为 null；数据库中存在 id 对应的任务；目标任务状态不为 "COMPLETED"'],
        ['后置条件', '返回更新后的 Task 对象；数据库中对应记录已更新；updateTime 由系统自动刷新'],
        ['异常条件',
         '(1) id 对应的任务不存在 → IllegalArgumentException("任务不存在")\n'
         '(2) 目标任务状态为 "COMPLETED" → IllegalArgumentException("已完成任务不允许修改")'],
        ['边界情况',
         '(1) 传入已完成任务的 id → 应被拒绝\n'
         '(2) 传入部分字段更新（如只更新 title）→ 应正常工作\n'
         '(3) id 为 null → NullPointerException'],
    ],
    col_widths=[2.5, 11.5],
)

# ---------- 方法 3: deleteTask ----------
add_heading_styled('3. void deleteTask(Long id) — 删除任务', level=2)

add_table(
    ['维度', '描述'],
    [
        ['前置条件', 'id 不为 null；数据库中存在 id 对应的任务'],
        ['后置条件', '数据库中对应记录被物理删除'],
        ['异常条件', 'id 对应的任务不存在 → IllegalArgumentException("任务不存在")'],
        ['边界情况',
         '(1) 重复删除同一个 id → 第二次应抛出异常\n'
         '(2) id 为 null → NullPointerException\n'
         '(3) 删除后调用 getTaskById 应返回 null'],
    ],
    col_widths=[2.5, 11.5],
)

# ---------- 方法 4: getTaskById ----------
add_heading_styled('4. Task getTaskById(Long id) — 按 ID 查询任务', level=2)

add_table(
    ['维度', '描述'],
    [
        ['前置条件', 'id 不为 null'],
        ['后置条件', '返回对应的 Task 对象；若不存在则返回 null'],
        ['异常条件', 'id 为 null → NullPointerException（取决于 mapper 实现）'],
        ['边界情况',
         '(1) 查询不存在的 id → 返回 null，不抛异常\n'
         '(2) 刚创建的任务 id → 应能查到完整信息'],
    ],
    col_widths=[2.5, 11.5],
)

# ---------- 方法 5: getAllTasks ----------
add_heading_styled('5. List<Task> getAllTasks() — 查询全部任务', level=2)

add_table(
    ['维度', '描述'],
    [
        ['前置条件', '无'],
        ['后置条件', '返回所有任务的列表；若数据库为空，返回空列表（不返回 null）'],
        ['异常条件', '无'],
        ['边界情况',
         '(1) 数据库无任务 → 返回空列表（size 为 0）\n'
         '(2) 数据库有任务 → 返回所有记录'],
    ],
    col_widths=[2.5, 11.5],
)

# ---------- 方法 6: getTasksByStatus ----------
add_heading_styled('6. List<Task> getTasksByStatus(String status) — 按状态筛选', level=2)

add_table(
    ['维度', '描述'],
    [
        ['前置条件', 'status 为有效值（"TODO" / "IN_PROGRESS" / "COMPLETED" 之一）'],
        ['后置条件', '返回所有状态为 status 的任务列表'],
        ['异常条件', '传入无效 status 值 → 返回空列表（由 mapper 的 SQL 查询决定）'],
        ['边界情况',
         '(1) 无匹配任务 → 返回空列表\n'
         '(2) 传入 null → 取决于实现（可能返回空列表或抛异常，建议返回空列表）\n'
         '(3) 传入空字符串 "" → 返回空列表'],
    ],
    col_widths=[2.5, 11.5],
)

# ---------- 方法 7: getTasksByPriority ----------
add_heading_styled('7. List<Task> getTasksByPriority(String priority) — 按优先级筛选', level=2)

add_table(
    ['维度', '描述'],
    [
        ['前置条件', 'priority 为有效值（"HIGH" / "MEDIUM" / "LOW" 之一）'],
        ['后置条件', '返回所有优先级为 priority 的任务列表'],
        ['异常条件', '传入无效 priority 值 → 返回空列表'],
        ['边界情况',
         '(1) 无匹配任务 → 返回空列表\n'
         '(2) 传入 null → 取决于实现\n'
         '(3) 传入空字符串 "" → 返回空列表'],
    ],
    col_widths=[2.5, 11.5],
)

# ---------- 方法 8: updateTaskStatus ----------
add_heading_styled('8. void updateTaskStatus(Long id, String status) — 更新任务状态', level=2)

add_table(
    ['维度', '描述'],
    [
        ['前置条件', 'id 不为 null；数据库中存在 id 对应的任务；status 为有效状态值'],
        ['后置条件', '对应任务的 status 被更新为新值；updateTime 被刷新'],
        ['异常条件',
         '(1) 任务不存在 → IllegalArgumentException("任务不存在")\n'
         '(2) 当前状态为 "COMPLETED" 且目标状态为 "TODO" → '
         'IllegalArgumentException("已完成任务不能回退到未开始状态")'],
        ['边界情况',
         '(1) "COMPLETED" → "IN_PROGRESS"：当前实现允许此操作（只禁止回退到 TODO）\n'
         '(2) 正常流转 "TODO" → "IN_PROGRESS" → "COMPLETED" 应正常工作\n'
         '(3) 空字符串作为 status → 取决于 mapper 实现\n'
         '(4) 状态未变化（如 TODO → TODO）→ 应正常执行'],
    ],
    col_widths=[2.5, 11.5],
)

# ---------- 方法 9: getTotalCount ----------
add_heading_styled('9. int getTotalCount() — 获取任务总数', level=2)

add_table(
    ['维度', '描述'],
    [
        ['前置条件', '无'],
        ['后置条件', '返回数据库中所有任务的总数'],
        ['异常条件', '无'],
        ['边界情况',
         '(1) 数据库为空 → 返回 0\n'
         '(2) 数据库有多条任务 → 返回正确数量'],
    ],
    col_widths=[2.5, 11.5],
)

# ---------- 方法 10: getCompletedCount ----------
add_heading_styled('10. int getCompletedCount() — 获取已完成任务数', level=2)

add_table(
    ['维度', '描述'],
    [
        ['前置条件', '无'],
        ['后置条件', '返回状态为 "COMPLETED" 的任务数量'],
        ['异常条件', '无'],
        ['边界情况',
         '(1) 无已完成任务 → 返回 0\n'
         '(2) 所有任务都已完成 → 返回值等于 getTotalCount()'],
    ],
    col_widths=[2.5, 11.5],
)

# ---------- 方法 11: getUnfinishedCount ----------
add_heading_styled('11. int getUnfinishedCount() — 获取未完成任务数', level=2)

add_table(
    ['维度', '描述'],
    [
        ['前置条件', '无'],
        ['后置条件', '返回 total - completed，即未完成任务数量'],
        ['异常条件', '无'],
        ['边界情况',
         '(1) 全部完成 → 返回 0\n'
         '(2) 全部未完成 → 返回值等于 getTotalCount()\n'
         '(3) 数据库为空 → 返回 0'],
    ],
    col_widths=[2.5, 11.5],
)

# ---------- 方法 12: getCompletionRate ----------
add_heading_styled('12. double getCompletionRate() — 获取完成率', level=2)

add_table(
    ['维度', '描述'],
    [
        ['前置条件', '无'],
        ['后置条件', '返回完成率百分比，计算公式为 (completed / total) × 100'],
        ['异常条件', '无'],
        ['边界情况',
         '(1) 任务总数为 0 → 返回 0.0（避免除零错误）\n'
         '(2) 全部完成 → 返回 100.0\n'
         '(3) 部分完成 → 返回 0~100 之间的值（如 3 个任务完成 1 个 → 33.33...）'],
    ],
    col_widths=[2.5, 11.5],
)

# 五、状态流转约束
add_heading_styled('五、状态流转约束', level=1)
add_para('根据需求文档中的业务不变量，任务状态的合法流转路径如下：')

state_flow = [
    '    ┌─────────────────────────────────────────────────┐',
    '    │                                                 │',
    '    │   TODO  ──────▶  IN_PROGRESS  ──────▶  COMPLETED│',
    '    │     ▲                              │      │     │',
    '    │     └────────── 禁止回退 ───────────┘      │     │',
    '    │                                   允许回退到     │',
    '    │                                   IN_PROGRESS   │',
    '    └─────────────────────────────────────────────────┘',
]
add_code_block(state_flow)
doc.add_paragraph()

add_para('约束规则：', bold=True)
rules = [
    'TODO → IN_PROGRESS：允许（正常开始任务）',
    'IN_PROGRESS → COMPLETED：允许（正常完成任务）',
    'COMPLETED → IN_PROGRESS：允许（当前实现未禁止此方向）',
    'COMPLETED → TODO：禁止（不允许将已完成任务回退到未开始状态）',
    'IN_PROGRESS → TODO：未在当前实现中明确限制，视为允许',
]
for i, r in enumerate(rules, 1):
    add_para(f'  {i}. {r}')

# 六、异常处理汇总
add_heading_styled('六、异常条件汇总', level=1)

add_table(
    ['方法', '异常条件', '异常类型', '错误信息'],
    [
        ['createTask', '任务标题为空', 'IllegalArgumentException', '任务标题不能为空'],
        ['createTask', '高优先级无截止时间', 'IllegalArgumentException', '高优先级任务必须设置截止时间'],
        ['createTask', '截止时间已过', 'IllegalArgumentException', '截止时间不能早于当前时间'],
        ['updateTask', '任务不存在', 'IllegalArgumentException', '任务不存在'],
        ['updateTask', '修改已完成任务', 'IllegalArgumentException', '已完成任务不允许修改'],
        ['deleteTask', '任务不存在', 'IllegalArgumentException', '任务不存在'],
        ['updateTaskStatus', '任务不存在', 'IllegalArgumentException', '任务不存在'],
        ['updateTaskStatus', '已完成回退到TODO', 'IllegalArgumentException', '已完成任务不能回退到未开始状态'],
    ],
    col_widths=[3, 3.5, 3.5, 4],
)

# 七、边界情况汇总
add_heading_styled('七、边界情况汇总', level=1)

add_table(
    ['场景', '输入', '期望行为'],
    [
        ['空标题', 'title = ""', '抛出 IllegalArgumentException'],
        ['纯空格标题', 'title = "   "', '抛出 IllegalArgumentException'],
        ['高优无截止时间', 'priority = "HIGH", deadline = null', '抛出 IllegalArgumentException'],
        ['过去截止时间', 'deadline = 2024-01-01', '抛出 IllegalArgumentException'],
        ['空数据库统计', '无任何任务', 'total = 0, completed = 0, rate = 0.0'],
        ['全部完成', '所有任务 status = COMPLETED', 'unfinishedCount = 0, rate = 100.0'],
        ['查询不存在的任务', 'id = 99999', '返回 null'],
        ['重复删除', '对同一 id 调用两次 deleteTask', '第二次抛出 IllegalArgumentException'],
        ['修改已完成任务', '对 COMPLETED 任务调用 updateTask', '抛出 IllegalArgumentException'],
        ['状态回退到 TODO', 'COMPLETED → TODO', '抛出 IllegalArgumentException'],
        ['空列表查询', 'getAllTasks() 时数据库无记录', '返回空列表（非 null）'],
        ['完成率除零', 'total = 0', '返回 0.0'],
    ],
    col_widths=[3, 5.5, 5.5],
)

# ============ 分页 ============
doc.add_page_break()

# ============ 附录：AI Prompt 记录 ============
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run('附录：AI Prompt 记录表')
run.bold = True
run.font.size = Pt(18)
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
doc.add_paragraph()

add_table(
    ['序号', 'Prompt 内容', 'AI 输出（摘要）', '人工修订'],
    [
        ['1',
         '给出 TaskService 接口签名和行为约束（前置条件、后置条件、异常条件、边界情况）',
         '输出 12 个方法的完整契约定义，包含前置/后置条件、异常处理、边界场景',
         '根据实际 TaskServiceImpl 代码校验：补充 "COMPLETED→IN_PROGRESS 允许" 的细节；确认除零保护逻辑'],
        ['2',
         '基于接口契约生成 JUnit 5 测试用例',
         'AI 将补全每个方法的正常路径、异常路径和边界条件测试代码',
         '待完成'],
    ],
    col_widths=[1.2, 4.5, 4.5, 3.8],
)

# ============ 页脚 ============
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— 全文完 —')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(128, 128, 128)

# ============ 保存 ============
output_path = '/Users/lyh/testclass/Software Construction/StudyFlow/实验二_接口签名与行为约束.docx'
doc.save(output_path)
print(f'Word 文档已生成：{output_path}')
